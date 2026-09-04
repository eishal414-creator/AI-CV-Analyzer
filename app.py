import os
import json
import re
from io import BytesIO

import streamlit as st
from pypdf import PdfReader
from docx import Document
from google import genai


# -----------------------------
# PAGE CONFIG
# -----------------------------
st.set_page_config(
    page_title="Resume ATS Analyzer",
    page_icon="📄",
    layout="wide",
)


# -----------------------------
# CUSTOM CSS
# -----------------------------
st.markdown(
    """
    <style>
    .main-title {
        font-size: 2.7rem;
        font-weight: 700;
        margin-bottom: 0.2rem;
    }

    .subtitle {
        color: #6b7280;
        font-size: 1.05rem;
        margin-bottom: 1.5rem;
    }

    .score-card {
        padding: 1.4rem;
        border: 1px solid rgba(128,128,128,0.25);
        border-radius: 14px;
        text-align: center;
        margin-bottom: 1rem;
    }

    .score-number {
        font-size: 3.2rem;
        font-weight: 800;
        margin: 0;
    }

    .small-note {
        color: #6b7280;
        font-size: 0.9rem;
    }
    </style>
    """,
    unsafe_allow_html=True,
)


# -----------------------------
# AI PROMPT
# -----------------------------
SYSTEM_PROMPT = """
You are an expert ATS resume reviewer, recruiter, and career coach.

Analyze the supplied resume carefully. Your job is to estimate how ATS-friendly
the resume is and provide practical improvements.

If a job description is supplied, evaluate the resume specifically against it.
If no job description is supplied, evaluate general ATS readiness.

Important:
- Do not invent information.
- Do not claim that your score is an official score from a real ATS.
- Base recommendations only on the supplied resume and job description.
- Be specific and actionable.
- Prefer standard ATS-friendly terminology.
- Identify missing or weak keywords only when they are relevant to the job
  description or the resume's apparent target role.

Return ONLY valid JSON using exactly this structure:

{
  "ats_score": 0,
  "score_breakdown": {
    "formatting": 0,
    "keywords": 0,
    "experience": 0,
    "skills": 0,
    "content": 0
  },
  "summary": "",
  "strengths": [],
  "improvements": [],
  "keyword_gaps": [],
  "formatting_issues": [],
  "action_plan": []
}

Scoring:
- ats_score: integer from 0 to 100.
- Each score_breakdown value: integer from 0 to 100.
- formatting: ATS parsing/readability and standard structure.
- keywords: relevant keywords and alignment.
- experience: clarity, relevance, achievements, and measurable impact.
- skills: technical/professional skills and their relevance.
- content: summary, education, projects, clarity, and completeness.

Keep each list concise:
- strengths: 3 to 6 items
- improvements: 4 to 8 items
- keyword_gaps: 3 to 10 items
- formatting_issues: 2 to 8 items
- action_plan: 4 to 8 items
"""


# -----------------------------
# FILE TEXT EXTRACTION
# -----------------------------
def extract_text(uploaded_file):
    """Extract text from PDF, DOCX, or TXT files."""
    data = uploaded_file.getvalue()
    filename = uploaded_file.name.lower()

    if filename.endswith(".pdf"):
        reader = PdfReader(BytesIO(data))
        pages = []

        for page in reader.pages:
            pages.append(page.extract_text() or "")

        return "\n".join(pages)

    if filename.endswith(".docx"):
        document = Document(BytesIO(data))

        paragraphs = [
            paragraph.text
            for paragraph in document.paragraphs
            if paragraph.text.strip()
        ]

        # Also extract text from tables because resumes sometimes use tables.
        for table in document.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                paragraphs.append(" | ".join(cells))

        return "\n".join(paragraphs)

    if filename.endswith(".txt"):
        return data.decode("utf-8", errors="ignore")

    raise ValueError(
        "Unsupported file type. Please upload a PDF, DOCX, or TXT file."
    )


# -----------------------------
# JSON CLEANING
# -----------------------------
def parse_json_response(text):
    """Convert Gemini's response into a Python dictionary."""
    if not text:
        raise ValueError("Gemini returned an empty response.")

    text = text.strip()

    # Remove markdown code fences if the model adds them.
    text = re.sub(r"^```(?:json)?\s*", "", text, flags=re.IGNORECASE)
    text = re.sub(r"\s*```$", "", text)

    try:
        return json.loads(text)
    except json.JSONDecodeError:
        # Try to recover a JSON object embedded in surrounding text.
        match = re.search(r"\{.*\}", text, flags=re.DOTALL)

        if not match:
            raise ValueError("The AI response was not valid JSON.")

        return json.loads(match.group(0))


# -----------------------------
# NORMALIZE AI RESULT
# -----------------------------
def normalize_result(result):
    """Make sure the expected fields exist and values are safe to display."""

    def safe_list(value):
        if isinstance(value, list):
            return [str(item) for item in value]
        return []

    breakdown = result.get("score_breakdown", {})

    if not isinstance(breakdown, dict):
        breakdown = {}

    score = result.get("ats_score", 0)

    try:
        score = int(score)
    except (TypeError, ValueError):
        score = 0

    score = max(0, min(100, score))

    normalized_breakdown = {}

    for key in ["formatting", "keywords", "experience", "skills", "content"]:
        try:
            value = int(breakdown.get(key, 0))
        except (TypeError, ValueError):
            value = 0

        normalized_breakdown[key] = max(0, min(100, value))

    return {
        "ats_score": score,
        "score_breakdown": normalized_breakdown,
        "summary": str(result.get("summary", "")),
        "strengths": safe_list(result.get("strengths", [])),
        "improvements": safe_list(result.get("improvements", [])),
        "keyword_gaps": safe_list(result.get("keyword_gaps", [])),
        "formatting_issues": safe_list(
            result.get("formatting_issues", [])
        ),
        "action_plan": safe_list(result.get("action_plan", [])),
    }


# -----------------------------
# API KEY
# -----------------------------
def get_api_key(user_key):
    """Get the API key from UI input, Streamlit Secrets, or environment."""
    if user_key and user_key.strip():
        return user_key.strip()

    try:
        secret_key = st.secrets.get("GEMINI_API_KEY", "")
        if secret_key:
            return secret_key
    except Exception:
        pass

    return os.getenv("GEMINI_API_KEY", "").strip()


# -----------------------------
# GEMINI ANALYSIS
# -----------------------------
def analyze_resume(resume_text, job_description, api_key, model_name):
    """Send resume text to Gemini and return structured analysis."""

    client = genai.Client(api_key=api_key)

    job_text = job_description.strip() if job_description else ""

    if not job_text:
        job_text = "No job description was provided."

    prompt = f"""
{SYSTEM_PROMPT}

====================
RESUME
====================

{resume_text}

====================
JOB DESCRIPTION
====================

{job_text}

====================
INSTRUCTIONS
====================

Analyze the resume now.

Return JSON only. Do not use markdown.
"""

    response = client.models.generate_content(
        model=model_name,
        contents=prompt,
        config={
            "temperature": 0.2,
            "response_mime_type": "application/json",
        },
    )

    return normalize_result(parse_json_response(response.text))


# -----------------------------
# HEADER
# -----------------------------
st.markdown(
    '<div class="main-title">📄 Resume ATS Analyzer</div>',
    unsafe_allow_html=True,
)

st.markdown(
    '<div class="subtitle">'
    "Upload your resume to estimate ATS readiness and get AI-powered "
    "improvement suggestions."
    "</div>",
    unsafe_allow_html=True,
)


# -----------------------------
# SIDEBAR
# -----------------------------
with st.sidebar:
    st.header("⚙️ Settings")

    api_key_input = st.text_input(
        "Gemini API Key",
        type="password",
        help=(
            "For local testing you can paste your API key here. "
            "For deployment, use Streamlit Secrets."
        ),
    )

    model_name = st.text_input(
        "Gemini model",
        value="gemini-2.5-flash",
    )

    st.divider()

    st.markdown("### 🎯 Optional")
    st.caption(
        "Paste a job description to get a more targeted ATS analysis."
    )

    job_description = st.text_area(
        "Job Description",
        height=260,
        placeholder=(
            "Paste the complete job description here..."
        ),
    )


# -----------------------------
# MAIN INPUT
# -----------------------------
uploaded_file = st.file_uploader(
    "📤 Upload your resume",
    type=["pdf", "docx", "txt"],
    help="Supported formats: PDF, DOCX, TXT",
)

if uploaded_file:
    st.success(
        f"Uploaded: **{uploaded_file.name}**"
    )


# -----------------------------
# ANALYZE
# -----------------------------
analyze_button = st.button(
    "🔍 Analyze Resume",
    type="primary",
    use_container_width=True,
    disabled=uploaded_file is None,
)


if analyze_button:

    api_key = get_api_key(api_key_input)

    if not api_key:
        st.error(
            "Gemini API key not found. Enter it in the sidebar or "
            "add GEMINI_API_KEY to Streamlit Secrets."
        )
        st.stop()

    try:
        # Extract resume
        with st.spinner("📖 Reading your resume..."):
            resume_text = extract_text(uploaded_file)

        if not resume_text.strip():
            st.error(
                "No readable text was found in this file. "
                "If your PDF is a scanned image, please use a text-based "
                "PDF or DOCX file."
            )
            st.stop()

        # Limit input size to keep requests practical.
        max_characters = 30000

        if len(resume_text) > max_characters:
            resume_text = resume_text[:max_characters]
            st.warning(
                "The resume was very long, so only the first "
                f"{max_characters:,} characters were analyzed."
            )

        # Analyze
        with st.spinner("🤖 Gemini is analyzing your resume..."):
            result = analyze_resume(
                resume_text=resume_text,
                job_description=job_description,
                api_key=api_key,
                model_name=model_name.strip() or "gemini-2.5-flash",
            )

        # Save result for download / rerender.
        st.session_state["analysis_result"] = result
        st.session_state["resume_name"] = uploaded_file.name

    except Exception as error:
        st.error(
            "Something went wrong while analyzing the resume."
        )

        st.exception(error)


# -----------------------------
# DISPLAY RESULT
# -----------------------------
if "analysis_result" in st.session_state:

    result = st.session_state["analysis_result"]
    score = result["ats_score"]

    st.divider()

    st.subheader("📊 ATS Readiness Score")

    score_col, details_col = st.columns([1, 2])

    with score_col:
        st.markdown(
            f"""
            <div class="score-card">
                <p class="small-note">Estimated ATS Score</p>
                <p class="score-number">{score}/100</p>
            </div>
            """,
            unsafe_allow_html=True,
        )

        st.progress(score / 100)

        if score >= 80:
            st.success("Strong ATS readiness")
        elif score >= 60:
            st.warning("Moderate ATS readiness")
        else:
            st.error("Needs improvement")

    with details_col:
        st.markdown("### 📈 Score Breakdown")

        breakdown = result["score_breakdown"]

        metric_cols = st.columns(5)

        metrics = [
            ("Formatting", breakdown["formatting"]),
            ("Keywords", breakdown["keywords"]),
            ("Experience", breakdown["experience"]),
            ("Skills", breakdown["skills"]),
            ("Content", breakdown["content"]),
        ]

        for column, (label, value) in zip(metric_cols, metrics):
            with column:
                st.metric(label, f"{value}/100")

    st.info(
        "This is an AI-based ATS readiness estimate, not an official "
        "score from a specific Applicant Tracking System."
    )

    # Summary
    st.subheader("📝 Overall Assessment")
    st.write(result["summary"])

    # Strengths / improvements
    left, right = st.columns(2)

    with left:
        st.markdown("### ✅ Strengths")

        if result["strengths"]:
            for item in result["strengths"]:
                st.write(f"• {item}")
        else:
            st.write("No major strengths were returned.")

    with right:
        st.markdown("### ⚠️ Improvements")

        if result["improvements"]:
            for item in result["improvements"]:
                st.write(f"• {item}")
        else:
            st.write("No specific improvements were returned.")

    # Keywords / formatting
    left, right = st.columns(2)

    with left:
        st.markdown("### 🔑 Keyword Gaps")

        if result["keyword_gaps"]:
            for item in result["keyword_gaps"]:
                st.write(f"• {item}")
        else:
            st.write("No major keyword gaps identified.")

    with right:
        st.markdown("### 🧩 Formatting Issues")

        if result["formatting_issues"]:
            for item in result["formatting_issues"]:
                st.write(f"• {item}")
        else:
            st.write("No major formatting issues identified.")

    # Action plan
    st.subheader("🚀 Recommended Action Plan")

    if result["action_plan"]:
        for index, item in enumerate(
            result["action_plan"],
            start=1,
        ):
            st.write(f"**{index}.** {item}")
    else:
        st.write("No action plan was returned.")

    # Download JSON
    st.divider()

    st.subheader("💾 Export Analysis")

    json_data = json.dumps(
        result,
        indent=2,
        ensure_ascii=False,
    )

    st.download_button(
        label="⬇️ Download Analysis as JSON",
        data=json_data,
        file_name="resume_ats_analysis.json",
        mime="application/json",
    )
